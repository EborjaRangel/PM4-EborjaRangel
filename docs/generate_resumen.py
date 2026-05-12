# -*- coding: utf-8 -*-
"""
Genera docs/RESUMEN_EJECUTIVO.docx con el resumen ejecutivo completo del
proyecto PULSE (entidades, rutas, páginas, formularios y flujos).

Requiere: python-docx
"""

from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


PULSE_BLUE = RGBColor(0x18, 0x77, 0xF2)
PULSE_DARK = RGBColor(0x1C, 0x1E, 0x21)
PULSE_GREY = RGBColor(0x65, 0x67, 0x6B)
HEADER_FILL = "1877F2"  # Azul Pulse para el header de tablas
ALT_ROW_FILL = "F0F6FF"  # Filas alternas suaves


# --------- Helpers ---------
def set_cell_shading(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def set_run_style(run, *, size=11, bold=False, italic=False, color=PULSE_DARK, font="Calibri"):
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Calibri"
        run.font.color.rgb = PULSE_BLUE if level <= 2 else PULSE_DARK


def add_paragraph(doc, text, *, bold=False, italic=False, size=11, color=PULSE_DARK, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_run_style(run, size=size, bold=bold, italic=italic, color=color)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    set_run_style(run)
    return p


def add_numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    run = p.add_run(text)
    set_run_style(run)
    return p


def add_table(doc: Document, headers: list[str], rows: list[list[str]], col_widths_cm: list[float] | None = None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.autofit = False

    # Header
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        set_run_style(run, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        set_cell_shading(cell, HEADER_FILL)

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.cell(r_idx + 1, c_idx)
            cell.text = ""
            run = cell.paragraphs[0].add_run(val)
            set_run_style(run)
            if r_idx % 2 == 1:
                set_cell_shading(cell, ALT_ROW_FILL)

    if col_widths_cm:
        for col_idx, w in enumerate(col_widths_cm):
            for row in table.rows:
                row.cells[col_idx].width = Cm(w)

    doc.add_paragraph()  # spacer


# --------- Documento ---------
def build_document(out_path: Path) -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ===== Portada =====
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_style(p.add_run("PULSE"), size=40, bold=True, color=PULSE_BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_style(
        p.add_run("Resumen ejecutivo del proyecto"),
        size=16, italic=True, color=PULSE_GREY,
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_style(
        p.add_run("Ecommerce premium · Next.js + Express + PostgreSQL"),
        size=12, color=PULSE_GREY,
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_style(p.add_run("Versión 11 de mayo de 2026"), size=11, color=PULSE_GREY)

    doc.add_paragraph()

    # ===== 1. ¿Qué es? =====
    add_heading(doc, "1. ¿Qué es?", level=1)
    add_paragraph(
        doc,
        "PULSE / Dreams Time es un e-commerce demostrativo de tecnología premium con dos "
        "aplicaciones independientes que pueden ejecutarse en local o exponerse al exterior con ngrok.",
    )
    add_table(
        doc,
        headers=["Capa", "Stack", "Puerto local", "Función"],
        rows=[
            ["Front", "Next.js 16 (App Router) + React 19 + Tailwind 4 + Leaflet", "3001", "Tienda, perfil de cliente y panel admin"],
            ["Back", "Express + TypeORM + PostgreSQL + JWT", "3000", "API REST que persiste todo en postgres://localhost:5432"],
        ],
        col_widths_cm=[1.8, 6.5, 2.0, 5.5],
    )
    add_paragraph(
        doc,
        "Para demos en móvil se levanta un único túnel ngrok hacia el frontend (puerto 3001); "
        "el front se comunica con el back vía proxy server-side, así no hace falta exponer la API.",
        italic=True, color=PULSE_GREY,
    )

    # ===== 2. Modelo de datos =====
    add_heading(doc, "2. Modelo de datos (PostgreSQL · TypeORM)", level=1)
    add_table(
        doc,
        headers=["Entidad", "Tabla", "Campos relevantes", "Para qué sirve"],
        rows=[
            ["User", "users", "id, name, email, address, phone, role, loginCount, lastLoginAt", "Cliente / admin registrado"],
            ["Credential", "credentials", "username, password (bcrypt)", "Credenciales del usuario"],
            ["Category", "categories", "id, name", "Categorías de productos"],
            ["Product", "products", "name, description, price, stock, image, images[], categoryId", "Catálogo de productos"],
            ["Order", "orders", "user, products[] (M:N)", "Orden histórica del back original (legacy)"],
            ["PurchaseRecord", "purchase_records", "clientUserId, userEmail, items, totals, shipping, status, createdAt", "Compra completada en el checkout simulado"],
            ["SavedCart", "saved_carts", "userId, items (json)", "Carrito persistido por usuario"],
            ["Address", "addresses", "id, userId, label, address, phone, lat, lng, isDefault", "Libreta de direcciones del cliente"],
        ],
        col_widths_cm=[2.5, 3.0, 6.0, 4.5],
    )
    add_paragraph(
        doc,
        "Al arrancar la API se ejecutan preloads: 9 categorías por defecto y, si no existe, "
        "un admin built-in para entrar al panel.",
        italic=True, color=PULSE_GREY,
    )

    # ===== 3. Backend — Rutas =====
    add_heading(doc, "3. Backend - Rutas (http://localhost:3000)", level=1)

    add_heading(doc, "3.1 Autenticación y usuarios (/users)", level=2)
    add_table(
        doc,
        headers=["Método", "Path", "Auth", "Para qué"],
        rows=[
            ["POST", "/users/register", "público", "Registra usuario + aviso WhatsApp opcional (CallMeBot)"],
            ["POST", "/users/login", "público", "Devuelve JWT + datos públicos del usuario"],
            ["GET", "/users/me", "JWT", "Datos del usuario actual"],
            ["PATCH", "/users/me", "JWT", "Actualiza address, phone o contraseña"],
            ["GET", "/users", "JWT admin", "Lista todos los usuarios"],
        ],
        col_widths_cm=[1.8, 4.5, 2.0, 7.7],
    )

    add_heading(doc, "3.2 Catálogo (/products y /categories)", level=2)
    add_table(
        doc,
        headers=["Método", "Path", "Auth", "Para qué"],
        rows=[
            ["GET", "/products", "público", "Catálogo completo"],
            ["GET", "/products/:id", "público", "Detalle de un producto"],
            ["POST", "/products", "público (demo)", "Crea o actualiza (pulseUpdateId)"],
            ["POST", "/products/:id/replace", "público", "Edición por reemplazo"],
            ["PATCH/PUT/DELETE", "/products/:id", "público", "Edita o elimina"],
            ["GET/POST/PUT/DELETE", "/products/categories[/:id]", "público", "CRUD de categorías (alias de /categories)"],
        ],
        col_widths_cm=[3.5, 5.5, 2.0, 5.0],
    )

    add_heading(doc, "3.3 Compras (/purchases)", level=2)
    add_table(
        doc,
        headers=["Método", "Path", "Auth", "Para qué"],
        rows=[
            ["POST", "/purchases", "público", "Registra una compra simulada"],
            ["GET", "/purchases/user/:clientUserId", "público", "Historial del usuario"],
            ["GET", "/purchases/:id", "público", "Endpoint que abre el QR del envío"],
        ],
        col_widths_cm=[1.8, 5.8, 2.0, 6.4],
    )

    add_heading(doc, "3.4 Direcciones (/addresses)", level=2)
    add_table(
        doc,
        headers=["Método", "Path", "Auth", "Para qué"],
        rows=[
            ["GET", "/addresses", "JWT", "Libreta del usuario actual"],
            ["POST", "/addresses", "JWT", "Crear dirección (marca default si es la 1ª)"],
            ["PUT", "/addresses/:id", "JWT", "Editar"],
            ["DELETE", "/addresses/:id", "JWT", "Eliminar (promueve otra como default)"],
            ["POST", "/addresses/:id/default", "JWT", "Marcar como predeterminada"],
        ],
        col_widths_cm=[1.8, 5.5, 2.0, 6.7],
    )

    add_heading(doc, "3.5 Otras", level=2)
    add_bullet(doc, "/orders - viejo CRUD de órdenes heredado del back original.")
    add_bullet(doc, "/cart - persistencia del carrito por usuario (SavedCart).")

    # ===== 4. Frontend — Páginas =====
    add_heading(doc, "4. Frontend - Resumen de cada página", level=1)

    add_heading(doc, "4.1 Públicas / cliente", level=2)
    add_table(
        doc,
        headers=["Ruta", "Archivo", "Qué hace"],
        rows=[
            ["/landing", "app/landing/page.tsx", "Hero con video, stats (+10K · 24/7 · 48h) y CTAs. Título: PULSE: ecommerce premium."],
            ["/home", "app/home/page.tsx", "Tienda principal con HomeCatalog + CardContainer. Filtra por nombre, descripción y categoría."],
            ["/product/[id]", "app/product/[id]/page.tsx", "Ficha de producto: imágenes, precio, stock, agregar al carrito."],
            ["/cart", "app/cart/page.tsx", "Carrito, totales, datos del cliente. Bloquea elegir dirección y pagar si no hay sesión."],
            ["/cart/ubicacion", "app/cart/ubicacion/page.tsx", "Picker de dirección (OSM + Nominatim + Geolocation). Muestra libreta y permite agregar otra."],
            ["/checkout", "app/checkout/page.tsx", "Pago simulado: detecta marca con logos SVG, valida Luhn, selector de dirección, Total en globo azul, POST /purchases."],
            ["/mis-compras", "app/mis-compras/page.tsx", "Historial del usuario: dirección, mapa, productos, totales y QR único hacia /envio/[id]."],
            ["/envio/[id]", "app/envio/[id]/page.tsx", "Página pública que abre el QR: dirección, teléfono, mapa, productos y totales. Footer sin enlaces."],
            ["/login", "app/login/page.tsx", "Login con email + password, devuelve JWT."],
            ["/register", "app/register/page.tsx", "Registro con RegisterForm (Formik + Yup)."],
            ["/profile", "app/profile/page.tsx", "Perfil: datos, ProfileContactForm, AddressBook con CRUD, cambio de contraseña."],
            ["/dashboard", "app/dashboard/page.tsx", "Vista de cliente con info resumida y accesos."],
        ],
        col_widths_cm=[3.0, 4.5, 8.5],
    )

    add_heading(doc, "4.2 Admin (/admin/products)", level=2)
    add_table(
        doc,
        headers=["Ruta", "Archivo", "Qué hace"],
        rows=[
            ["/admin/products", "app/admin/products/page.tsx", "Hub con links a las acciones del admin."],
            ["/admin/products/alta", "app/admin/products/alta/page.tsx", "Form de alta con selector de categoría (DB) y hasta 5 imágenes."],
            ["/admin/products/editar", "app/admin/products/editar/page.tsx", "Edición de productos, incluida la categoría."],
            ["/admin/products/baja", "app/admin/products/baja/page.tsx", "Borra productos del catálogo."],
            ["/admin/products/categorias", "app/admin/products/categorias/page.tsx", "CRUD de categorías (persistente en Postgres)."],
        ],
        col_widths_cm=[4.0, 5.5, 6.5],
    )

    # ===== 5. Formularios =====
    add_heading(doc, "5. Formularios - qué validan y qué hacen", level=1)
    add_table(
        doc,
        headers=["Formulario", "Ubicación", "Campos", "Acción al enviar"],
        rows=[
            ["Registro", "RegisterForm.tsx", "name, email, password, confirm, address, phone", "POST /users/register; si OK redirige a /login"],
            ["Login", "/login", "email, password", "POST /users/login; guarda JWT en localStorage y redirige a home"],
            ["Perfil - contacto", "ProfileContactForm.tsx", "address, phone", "PATCH /users/me"],
            ["Perfil - password", "/profile", "currentPassword, newPassword, confirm", "PATCH /users/me con cambio de contraseña"],
            ["Dirección de envío", "OpenStreetMapDeliveryPicker.tsx", "label, calle, exterior, colonia, CP, ciudad, entre A/B, teléfono, isDefault", "POST /addresses (libreta) + PATCH /users/me como fallback"],
            ["Editar dirección", "AddressBook.tsx (inline)", "label, address, phone", "PUT /addresses/:id"],
            ["Alta de producto", "/admin/products/alta", "name, description, price, stock, categoryId, hasta 5 imágenes", "POST /products"],
            ["Editar producto", "/admin/products/editar", "mismos campos + id", "POST /products con pulseUpdateId"],
            ["Alta de categoría", "/admin/products/categorias", "name", "POST /products/categories"],
            ["Editar categoría", "/admin/products/categorias", "id, name", "PUT /products/categories/:id"],
            ["Checkout (tarjeta)", "/checkout", "holder, PAN, exp, CVV, dirección seleccionada", "Valida Luhn, arma snapshot, POST /purchases, vacía carrito"],
            ["Buscador del catálogo", "CardContainer.tsx", "query", "Filtra en memoria por nombre / descripción / categoría"],
            ["Buscador CP / dirección", "Mapa /cart/ubicacion", "cpQuery, addressQuery", "GET /api/geocode/search (Nominatim)"],
        ],
        col_widths_cm=[3.5, 3.5, 5.0, 4.0],
    )

    # ===== 6. Flujos clave =====
    add_heading(doc, "6. Flujos clave (de punta a punta)", level=1)

    add_heading(doc, "A) Registro → primera compra", level=2)
    for step in [
        "/register crea usuario en users + credentials.",
        "/login devuelve un JWT que se guarda en localStorage.",
        "/home muestra el catálogo desde GET /products; el buscador filtra por nombre, descripción y categoría.",
        "/product/[id] permite agregar al carrito (localStorage + SavedCart si hay sesión).",
        "/cart revisa cantidades y totales.",
        "/cart/ubicacion crea la primera dirección con label 'Casa' y la marca como default.",
        "/checkout elige dirección (radio cards) + tarjeta simulada; se guarda el PurchaseRecord con dirección y lat/lng.",
        "/mis-compras lista el pedido con su QR único.",
        "Al escanear el QR en el celular se abre /envio/<id> con los datos completos para el repartidor.",
    ]:
        add_numbered(doc, step)

    add_heading(doc, "B) Admin gestiona catálogo", level=2)
    for step in [
        "Login con el admin built-in (mostrado en cada página de admin).",
        "/admin/products/categorias crea una categoría nueva (persiste en Postgres).",
        "/admin/products/alta da de alta un producto con esa categoría.",
        "/admin/products/editar permite cambiar precio, stock, imagen y categoría.",
        "/admin/products/baja elimina el producto del catálogo.",
    ]:
        add_numbered(doc, step)

    add_heading(doc, "C) Cliente con varias direcciones", level=2)
    for step in [
        "En /profile o /cart/ubicacion ve su AddressBook con las direcciones guardadas.",
        "Puede editar, eliminar o marcar como predeterminada cualquier dirección.",
        "En /checkout aparecen como radio cards; la predeterminada queda preseleccionada.",
        "Al pagar, el PurchaseRecord guarda esa dirección con sus coordenadas.",
    ]:
        add_numbered(doc, step)

    # ===== 7. Infraestructura =====
    add_heading(doc, "7. Infraestructura, despliegue y dev", level=1)
    add_bullet(doc, "Local solo PC: Postgres encendido + 'npm run dev' en back y front.")
    add_bullet(doc, "Compartir con el celular (ngrok): 'ngrok http 3001' → URL pública del front; el back se llama vía proxy server-side (/pulse-api-proxy o /api/pulse-backend).")
    add_paragraph(doc, "Variables de entorno:", bold=True)
    add_bullet(doc, "back/.env: DB_*, JWT_SECRET, WHATSAPP_CALLMEBOT_APIKEY (opcional).")
    add_bullet(doc, "front/.env.local: NEXT_PUBLIC_GOOGLE_MAPS_API_KEY (opcional), NEXT_PUBLIC_API_URL / NEXT_PUBLIC_USE_LOCAL_API (modo ngrok).")
    add_paragraph(doc, "Despliegue:", bold=True)
    add_bullet(doc, "Back en Render con Postgres administrado.")
    add_bullet(doc, "Front en Vercel con PULSE_BACKEND_URL apuntando al Render.")
    add_bullet(doc, "Alternativa: front en Firebase App Hosting (ver FIREBASE.md).")

    # ===== 8. Funcionalidades destacadas =====
    add_heading(doc, "8. Funcionalidades destacadas", level=1)
    add_table(
        doc,
        headers=["Feature", "Dónde vive", "Resumen"],
        rows=[
            ["Categorías persistentes", "back Category + admin/categorias", "CRUD desde la UI, guardado en Postgres con synchronize: true."],
            ["Libreta de direcciones múltiples", "back Address + AddressBook + checkout selector", "Cada cliente tiene N direcciones, una default; al pagar elige una."],
            ["QR por pedido", "PurchaseQrCode + /envio/[id]", "qrcode.react genera SVG cuyo contenido es <origin>/envio/<id>; footer reducido."],
            ["Geolocalización en el mapa", "OpenStreetMapDeliveryPicker", "Solo se activa si el permiso ya está concedido; nunca lo pide de forma intrusiva."],
            ["Mapa sin Google", "Leaflet + Nominatim", "Modo fallback gratuito sin API key; con la key de Google cambia a Google Maps."],
            ["Total destacado", "Checkout", "Globo azul con degradado, sombra y 'TOTAL' en azul Pulse."],
            ["Logos de tarjeta", "Checkout", "SVG inline para Visa, Mastercard y Amex al detectar el PAN."],
            ["Búsqueda enriquecida", "/home", "Match por nombre / descripción / categoría con normalización de acentos."],
            ["Sesión obligatoria en el carrito", "/cart", "Botones de dirección y pago deshabilitados si no hay sesión."],
        ],
        col_widths_cm=[4.0, 4.5, 7.5],
    )

    # Cierre
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_style(
        p.add_run("Documento generado automáticamente desde el repositorio PULSE."),
        size=10, italic=True, color=PULSE_GREY,
    )

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


if __name__ == "__main__":
    out = Path(__file__).resolve().parent / "RESUMEN_EJECUTIVO.docx"
    build_document(out)
    print(f"OK -> {out}")
